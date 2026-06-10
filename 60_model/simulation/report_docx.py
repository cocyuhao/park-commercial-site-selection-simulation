from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


FORBIDDEN_HUMAN_TEXT = [
    "needs_review",
    "not_final",
    "output_status",
    "validation_status",
    "debug",
    "payload",
    "traceback",
    "ConnectError",
    "external" + "_preview_only",
    "API contract",
    "smoke test",
    "DeepSeek",
    "DOCX text is planning material",
    "measured operation data",
]


PUBLIC_TEXT_REPLACEMENTS = {
    "validation_status": "证据状态",
    "DeepSeek": "模型初稿",
    "DOCX text is planning material, not final measured operation data.": "策划资料属于方案材料，不是最终实测运营数据。",
    "Numbers or claims still need evidence catalog linkage before final recommendations.": "数字和判断进入最终建议前仍需完成证据目录关联。",
}


def _humanize_text(text: str) -> str:
    for old, new in PUBLIC_TEXT_REPLACEMENTS.items():
        text = text.replace(old, new)
    return text


def _text(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    text = _humanize_text(str(value).strip())
    return text or fallback


def _brief(value: Any, limit: int = 90) -> str:
    text = _text(value)
    if len(text) <= limit:
        return text
    return text[:limit].rstrip("，。；; ") + "..."


def _items(value: Any) -> list[str]:
    if isinstance(value, list):
        return [_text(item) for item in value if _text(item)]
    text = _text(value)
    if not text:
        return []
    parts = re.split(r"[；;]\s*|\n+", text)
    return [part.strip(" -") for part in parts if part.strip(" -")]


def _calibration_layer_label(value: Any) -> str:
    labels = {
        "official_macro_boundary": "官方宏观边界",
        "local_bigdata_profile": "本地大数据画像",
        "local_device_price_proxy": "设备价格代理",
        "local_poi_price_signal": "竞品价格线索",
        "local_poi_demand_signal": "本地需求热度线索",
        "plan_assumption_needs_review": "方案假设待复核",
        "local_user_supplement": "用户补充校准输入",
    }
    return labels.get(str(value or ""), str(value or "待分层"))


def _readiness_groups(report: dict[str, Any]) -> dict[str, list[str]]:
    readiness = report.get("simulation_readiness", {})
    return {
        "can_do": _items(readiness.get("can_do") or readiness.get("can_run_now")),
        "cannot_claim": _items(readiness.get("cannot_claim") or readiness.get("cannot_claim_yet")),
        "required_inputs": _items(readiness.get("required_inputs") or readiness.get("blocking_inputs")),
    }


def _set_east_asia_font(run: Any, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell: Any, text: Any, *, bold: bool = False, color: str = "111827") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(_text(text))
    _set_east_asia_font(run)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_table_borders(table: Any, color: str = "D1D5DB") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(5)


def _paragraph(doc: Document, text: Any, *, bold: bool = False, color: str = "111827", size: float = 10.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(_text(text))
    _set_east_asia_font(run)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _bullet(doc: Document, text: Any) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.63)
    paragraph.paragraph_format.first_line_indent = Cm(-0.21)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(_text(text))
    _set_east_asia_font(run)
    run.font.size = Pt(10.2)


def _numbered(doc: Document, text: Any) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Cm(0.63)
    paragraph.paragraph_format.first_line_indent = Cm(-0.21)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(_text(text))
    _set_east_asia_font(run)
    run.font.size = Pt(10.2)


def _small_table(doc: Document, rows: list[list[Any]], widths: list[float] | None = None) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_table_borders(table)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if widths and col_idx < len(widths):
                cell.width = Cm(widths[col_idx])
            _set_cell_text(cell, value, bold=row_idx == 0)
            if row_idx == 0:
                _shade_cell(cell, "F2F4F7")
    doc.add_paragraph()


def _add_cover(doc: Document, report: dict[str, Any]) -> None:
    title = _text(report.get("title"), "奥森商业改造综合评估与修正建议工作稿")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    _set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    _paragraph(doc, "商业选址与运营仿真工作稿", bold=True, color="2E74B5", size=12)
    _paragraph(doc, "本文件用于内部研判、方案修正和定案前复核；进入客户正式汇报前需完成 CAD 控制点校准、现场/运营复核、财务参数复核和人工审阅。", color="374151")

    meta_rows = [
        ["项目对象", "北京奥林匹克森林公园商业改造节点"],
        ["报告版本", f"工作稿 / {datetime.now().strftime('%Y-%m-%d')}"],
        ["交付格式", "DOCX 正式阅读稿；Markdown/JSON 仅作内部版本和验证记录"],
        ["核心边界", "不声明最终 ROI、最终收益、最终排名或完整仿真结果"],
    ]
    _small_table(doc, meta_rows, [3.2, 12.0])


def _add_summary(doc: Document, report: dict[str, Any]) -> None:
    doc.add_heading("1. 执行摘要", level=1)
    summary_items = _items(report.get("executive_summary") or report.get("summary"))[:5]
    for item in summary_items:
        _bullet(doc, item)

    doc.add_heading("一页式判断", level=2)
    readiness = _readiness_groups(report)
    rows = [
        ["现在可以做", "现在不能宣称", "定案前复核项"],
        [
            "\n".join(readiness["can_do"][:4]),
            "\n".join(readiness["cannot_claim"][:4]),
            "\n".join(readiness["required_inputs"][:5]),
        ],
    ]
    _small_table(doc, rows, [5.0, 5.0, 5.0])


def _add_evidence(doc: Document, report: dict[str, Any]) -> None:
    doc.add_heading("2. 关键依据", level=1)
    foundation = report.get("source_foundation", {})
    evidence = foundation.get("evidence_highlights", [])[:8]
    rows = [["依据", "数值/结论", "来源"]]
    for row in evidence:
        if isinstance(row, dict):
            rows.append([
                row.get("indicator_name", "证据"),
                f"{row.get('value', '')}{row.get('unit', '')}",
                row.get("source_file", ""),
            ])
        else:
            rows.append(["指标摘要", row, "证据台账 / 已抽取资料"])
    _small_table(doc, rows, [5.0, 3.2, 7.0])

    cad = foundation.get("cad", {})
    doc.add_heading("CAD 与图纸处理", level=2)
    _paragraph(doc, cad.get("boundary_note", "CAD 已用于识别图纸锚点；进入路径级仿真前仍需做控制点校准。"), color="374151")
    anchor_rows = [["图纸锚点", "图层", "CAD 坐标", "说明"]]
    for item in cad.get("keyword_anchors", [])[:7]:
        anchor_rows.append([
            item.get("label") or item.get("keyword"),
            item.get("layer", ""),
            f"{item.get('x', '')}, {item.get('y', '')}",
            "仅作图纸锚点；待控制点校准后进入 GIS",
        ])
    _small_table(doc, anchor_rows, [4.0, 2.2, 4.2, 5.2])


def _add_expert_basis(doc: Document, report: dict[str, Any]) -> None:
    basis = report.get("expert_review_basis") or {}
    if not basis:
        return
    doc.add_heading("3. 专家评审底座", level=1)
    screened = basis.get("screened_count") or 0
    completed = basis.get("completed_query_count") or 0
    query_total = basis.get("query_total") or 0
    if screened:
        _paragraph(
            doc,
            f"本轮已完成 {completed}/{query_total} 组主题检索，筛出 {screened} 条高相关研究候选。研究底座用于约束评审维度和证据门槛，不直接替代奥森的现场数据、审批意见或试运营结果。",
            color="374151",
        )
    beijing = basis.get("beijing_context") or {}
    if beijing.get("items"):
        doc.add_heading("收入与消费上位边界", level=2)
        for item in beijing.get("items", [])[:4]:
            _bullet(doc, item)
        if beijing.get("use_boundary"):
            _paragraph(doc, beijing["use_boundary"], bold=True, color="9A3412")
    dimensions = basis.get("dimensions") or []
    if dimensions:
        doc.add_heading("节点评审维度", level=2)
        rows = [["维度", "检查重点"]]
        for row in dimensions[:15]:
            if isinstance(row, list) and len(row) >= 2:
                rows.append([row[0], row[1]])
        _small_table(doc, rows, [4.0, 11.5])


def _add_real_calibration_context(doc: Document, report: dict[str, Any]) -> None:
    context = report.get("real_calibration_context") or {}
    doc.add_heading("4. 真实校准输入与使用边界", level=1)
    if not context.get("count"):
        _paragraph(doc, "当前还没有可复跑真实校准输入；报告只能保留为方法工作稿，不能进入收益或排名判断。", color="9A3412", bold=True)
        return

    _paragraph(
        doc,
        f"本轮已接入 {context.get('count')} 条真实校准输入，用于约束收入/消费边界、本地画像、设备价格代理、竞品价格和方案假设。它们必须分层使用，不能混成最终收益、排名或投资结论。",
        color="374151",
    )
    if context.get("report_rule"):
        _paragraph(doc, f"使用规则：{context['report_rule']}", bold=True, color="9A3412")

    strengths = context.get("source_strength_counts") or {}
    if strengths:
        layer_rows = [["分层", "数量"]]
        for layer, count in strengths.items():
            if count:
                layer_rows.append([_calibration_layer_label(layer), f"{count} 条"])
        _small_table(doc, layer_rows, [5.2, 3.0])

    rows = [["输入编号", "分层", "指标", "数值", "进入仿真的用法", "不能直接宣称"]]
    for item in (context.get("items") or [])[:10]:
        value = f"{_text(item.get('value'))}{_text(item.get('unit'))}"
        rows.append(
            [
                item.get("calibration_id", ""),
                _calibration_layer_label(item.get("source_strength")),
                item.get("indicator_name", ""),
                value,
                _brief(item.get("simulation_use"), 60),
                _brief(item.get("cannot_claim"), 60),
            ]
        )
    _small_table(doc, rows, [1.8, 2.4, 2.8, 2.2, 3.4, 3.2])

    missing = _items(context.get("missing_before_final"))[:6]
    if missing:
        doc.add_heading("进入定案前内部复核项", level=2)
        for item in missing:
            _bullet(doc, item)


def _add_feature_scene_context(doc: Document, report: dict[str, Any]) -> None:
    context = report.get("controlled_feature_scene_context") or {}
    doc.add_heading("5. 人物场景输入与收入价格带", level=1)
    if context.get("count"):
        _paragraph(
            doc,
            f"本轮已有 {context.get('count')} 条用户采用或锁定的人物场景进入报告输入。它们用于讨论收入水平、消费价格带、时段天气和节点动作的敏感性，不代表真实客群占比或最终仿真结果。",
            color="374151",
        )
        rows = [["编号", "场景", "收入/价格带", "时段/天气/空间", "建议动作", "复核证据"]]
        for item in (context.get("items") or [])[:8]:
            rows.append(
                [
                    item.get("derivative_id", ""),
                    item.get("title", ""),
                    " / ".join(
                        value
                        for value in [
                            _text(item.get("income_segment_name")),
                            _text(item.get("income_price_band")),
                        ]
                        if value
                    ),
                    "；".join(
                        value
                        for value in [
                            _text(item.get("time_band_name")),
                            _text(item.get("weather_name")),
                            _text(item.get("node_context_name")),
                        ]
                        if value
                    ),
                    item.get("candidate_supply_action_name", ""),
                    item.get("data_needed", ""),
                ]
            )
        _small_table(doc, rows, [1.9, 2.9, 3.0, 3.3, 2.4, 2.0])
    else:
        _paragraph(
            doc,
            context.get("empty_state", "当前还没有采用或锁定的人物场景；报告只能引用覆盖池作为方法底座，不能声明已完成客群仿真。"),
            color="9A3412",
            bold=True,
        )
    if context.get("report_rule"):
        _paragraph(doc, f"使用规则：{context['report_rule']}", color="374151")


def _add_nodes(doc: Document, report: dict[str, Any]) -> None:
    doc.add_heading("6. 六个节点判断与修改建议", level=1)
    overview_rows = [["节点", "推进判断", "第一动作"]]
    for node in report.get("nodes", []):
        review = node.get("implementation_review") or {}
        overview_rows.append([
            node.get("node_name", ""),
            _brief(node.get("priority_stage", ""), 34),
            _brief(review.get("recommended_path") or node.get("improvement", ""), 72),
        ])
    _small_table(doc, overview_rows, [3.5, 3.2, 9.0])

    for idx, node in enumerate(report.get("nodes", []), start=1):
        review = node.get("implementation_review") or {}
        doc.add_heading(f"6.{idx} {node.get('node_name', '')}", level=2)
        _paragraph(doc, f"推进类型：{node.get('priority_stage', '待判断')}", bold=True, color="1F4D78")
        if node.get("original_plan"):
            _paragraph(doc, f"原计划理解：{node.get('original_plan')}")
        if node.get("improvement"):
            _paragraph(doc, f"修正建议：{node.get('improvement')}", bold=True)
        if review:
            rows = [["检查项", "当前判断"]]
            rows.append(["服务对象", "；".join(_items(review.get("target_segments"))[:5])])
            rows.append(["需求触发", "；".join(_items(review.get("demand_triggers"))[:5])])
            rows.append(["收入与价格带", "；".join(_items(review.get("income_and_price_band"))[:3])])
            rows.append(["推荐路径", review.get("recommended_path", "")])
            _small_table(doc, rows, [3.2, 12.3])

            options = review.get("options") or []
            if options:
                doc.add_heading("方案比较", level=3)
                option_rows = [["方案", "怎么做", "适用条件", "先决条件", "主要风险"]]
                for option in options[:3]:
                    option_rows.append([
                        option.get("name", ""),
                        _brief(option.get("what_to_do", ""), 72),
                        _brief(option.get("best_for", ""), 44),
                        "；".join(_items(option.get("prerequisites"))[:3]),
                        "；".join(_items(option.get("risks"))[:3]),
                    ])
                _small_table(doc, option_rows, [2.1, 4.0, 3.0, 3.6, 3.2])

            for title, key in [
                ("时间、天气与季节", "time_weather"),
                ("周边复核口径", "surrounding_context_needed"),
                ("风险控制", "risk_controls"),
                ("哪些证据会改变判断", "evidence_that_changes_decision"),
            ]:
                values = _items(review.get(key))[:5]
                if values:
                    doc.add_heading(title, level=3)
                    for item in values:
                        _bullet(doc, item)
        required = _items(node.get("required_inputs"))[:8]
        if required:
            _paragraph(doc, "定案前需要复核：", bold=True, color="374151")
            for item in required:
                _bullet(doc, item)


def _add_recommendations(doc: Document, report: dict[str, Any]) -> None:
    doc.add_heading("7. 综合修改意见", level=1)
    for item in _items(report.get("revision_advice")):
        _bullet(doc, item)

    doc.add_heading("8. 仿真与定案边界", level=1)
    readiness = _readiness_groups(report)
    for title, key in [("现在可以做", "can_do"), ("现在不能宣称", "cannot_claim"), ("定案前复核项", "required_inputs")]:
        doc.add_heading(title, level=2)
        for item in readiness[key]:
            _bullet(doc, item)

    doc.add_heading("9. 当前推进事项", level=1)
    for item in _items(report.get("next_actions")):
        _numbered(doc, item)


def _add_appendix(doc: Document, report: dict[str, Any]) -> None:
    doc.add_heading("附录：使用边界与资料来源", level=1)
    _paragraph(doc, "本稿用于内部研判、方案修正和定案前复核。进入客户正式汇报前，应完成 CAD 控制点校准、现场/运营复核、财务参数复核和人工审阅。")
    foundation = report.get("source_foundation", {})

    method_basis = _items(report.get("method_basis"))
    if method_basis:
        doc.add_heading("方法依据", level=2)
        for item in method_basis[:5]:
            _bullet(doc, item)

    current_judgements = _items(report.get("current_judgements"))
    if current_judgements:
        doc.add_heading("现阶段判断边界", level=2)
        for item in current_judgements[:5]:
            _bullet(doc, item)

    assets = _items(foundation.get("assets"))
    if assets:
        doc.add_heading("已纳入的资料资产", level=2)
        for item in assets[:8]:
            _bullet(doc, item)

    plan_profile = foundation.get("plan_profile", {})
    if isinstance(plan_profile, dict) and plan_profile:
        doc.add_heading("策划资料读取口径", level=2)
        rows = [["项目", "结果"]]
        rows.append(["段落数量", plan_profile.get("paragraph_count", "")])
        rows.append(["表格数量", plan_profile.get("table_count", "")])
        keyword_counts = plan_profile.get("keyword_counts", {})
        if isinstance(keyword_counts, dict):
            keyword_summary = "；".join(f"{key} {value}" for key, value in list(keyword_counts.items())[:8])
            rows.append(["关键词命中", keyword_summary])
        limitations = plan_profile.get("limitations", [])
        if limitations:
            rows.append(["边界说明", "；".join(_items(limitations)[:3])])
        _small_table(doc, rows, [3.5, 11.5])


def _audit_text(docx_path: Path) -> dict[str, Any]:
    import zipfile
    import xml.etree.ElementTree as ET

    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    texts = [node.text or "" for node in root.iter() if node.tag.endswith("}t")]
    full_text = "\n".join(texts)
    return {
        "bytes": docx_path.stat().st_size,
        "text_length": len(full_text),
        "forbidden": [word for word in FORBIDDEN_HUMAN_TEXT if word in full_text],
        "required": {
            term: term in full_text
            for term in [
                "执行摘要",
                "关键依据",
                "专家评审底座",
                "真实校准输入",
                "官方宏观边界",
                "人物场景输入",
                "收入/价格带",
                "六个节点判断与修改建议",
                "综合修改意见",
                "当前推进事项",
                "控制点校准",
                "周边人口与收入",
                "桃花源白房子",
                "南门地下预埋空间",
            ]
        },
    }


def write_site_selection_docx(report: dict[str, Any], out_dir: Path) -> dict[str, Any]:
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "osen_integrated_site_selection_report_20260606.docx"
    audit_path = out_dir.parent / "40_quality_evidence" / "osen_integrated_report_docx_audit_20260606.json"

    doc = Document()
    _style_document(doc)
    _add_cover(doc, report)
    _add_summary(doc, report)
    _add_evidence(doc, report)
    _add_expert_basis(doc, report)
    _add_real_calibration_context(doc, report)
    _add_feature_scene_context(doc, report)
    _add_nodes(doc, report)
    _add_recommendations(doc, report)
    _add_appendix(doc, report)
    doc.save(docx_path)

    audit = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "docx_path": str(docx_path),
        "writing_format_research": "10_research/report_docx_writing_format_research_20260606.md",
        "design_preset": "standard_business_brief with Chinese A4 delivery override",
        "status": "pass",
        "structural_audit": _audit_text(docx_path),
    }
    if audit["structural_audit"]["forbidden"] or not all(audit["structural_audit"]["required"].values()):
        audit["status"] = "fail"
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if audit["status"] != "pass":
        raise RuntimeError(f"DOCX audit failed: {audit}")
    return {
        "docx": str(docx_path),
        "audit": str(audit_path),
        "byte_size": docx_path.stat().st_size,
    }
