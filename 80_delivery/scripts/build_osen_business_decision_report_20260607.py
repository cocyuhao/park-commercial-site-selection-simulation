from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from build_osen_prediction_adjustment_report_20260607 import build_basis  # noqa: E402


OUT_DIR = ROOT / "80_delivery"
QUALITY_DIR = ROOT / "40_quality_evidence"
STATIC_DIR = ROOT / "90_p6_expert_dashboard" / "static"

DOCX_OUT = OUT_DIR / "osen_business_decision_report_20260607.docx"
MD_OUT = OUT_DIR / "osen_business_decision_report_20260607.md"
BASIS_OUT = QUALITY_DIR / "osen_business_decision_report_basis_20260607.json"
EVIDENCE_MAP_OUT = QUALITY_DIR / "osen_platform_report_evidence_map_20260607.json"
STATIC_DOCX = STATIC_DIR / DOCX_OUT.name
STATIC_BASIS = STATIC_DIR / BASIS_OUT.name

FONT = "Microsoft YaHei"
FONT_LATIN = "Aptos"
INK = "18252D"
MUTED = "5E6B73"
BLUE = "173E63"
BLUE_2 = "245D7A"
GREEN = "0B5A46"
GOLD = "7B5A00"
RED = "8A2F20"
PAPER = "F7F9FA"
LIGHT_BLUE = "EAF2F7"
LIGHT_GREEN = "EAF5F0"
LIGHT_GOLD = "F8F1D9"
LIGHT_RED = "F8EDEA"
WHITE = "FFFFFF"


NODE_DECISIONS = [
    {
        "name": "南门露天剧场",
        "area": "草坪/湖畔空间",
        "priority": "第一批主推试点",
        "original": "湖畔阳光烘焙工坊 + 四季草坪生活场，含咖啡烘焙、市集、音乐、亲子工坊。",
        "forecast": "与现有客流和偏好最匹配。咖啡、冷饮、美食、亲子和草坪停留都能被现有数据解释，适合作为首个样板。",
        "adjustment": "用“烘焙 + 咖啡/冷饮 + 草坪活动 + 移动补给”组合切入；演出只做低噪声、活动日组件。",
        "operating_moves": [
            "周末与节假日设置咖啡/冷饮移动点、烘焙轻餐和草坪活动包。",
            "下午、活动前后、亲子停留时段优先排班，高温/雨天切到室内轻活动。",
            "把座椅、遮阴、排队动线和垃圾回收作为商业动作的一部分，不当成后勤附属。",
        ],
        "price": "咖啡/冷饮 20-45 元；烘焙轻餐 35-98 元；草坪套餐/亲子活动 68-168 元。",
        "evidence": ["峰值客流足够", "咖啡 TGI 286", "冷饮 TGI 332", "亲子/票务偏好高"],
        "boundary": "不以高噪声演出、重舞台或夜间餐酒作为首期收入主假设。",
    },
    {
        "name": "桃花源白房子",
        "area": "155㎡",
        "priority": "第一批低改造试点",
        "original": "活动+轻餐，多功能活动据点，承接婚礼/摄影、亲子/团建、咖啡茶饮和季节下午茶。",
        "forecast": "花海、拍照、轻餐和预约活动逻辑成立，但 155㎡ 空间决定它更适合做“轻量高周转”，不是重厨房。",
        "adjustment": "保留“白房子 + 花海 + 轻餐”定位，先做可移动外摆、预约活动、咖啡/冷饮和花季小活动。",
        "operating_moves": [
            "花季与周末主推摄影预约、生日小包场和冷餐会。",
            "平日保留咖啡、茶饮、轻食和后台活动预约，避免空间空置。",
            "通过套餐与预约制控制客流，不用满堂散客模型压垮现场。",
        ],
        "price": "日常 20-60 元；预约活动/冷餐会按 80-180 元/人做情景测算。",
        "evidence": ["CAD/PDF 图纸识别到白房子 155㎡", "咖啡/冷饮偏好强", "亲子与活动型需求可解释"],
        "boundary": "首期不做重厨房、大体量阳光房和高固定成本装修。",
    },
    {
        "name": "12#西分区管理中心",
        "area": "1059.39㎡",
        "priority": "第一批品牌联名试点",
        "original": "方案一 Live House；方案二北冰洋/义利国潮轻餐、文创和亲子体验。",
        "forecast": "北冰洋/义利式白天品牌联名更稳。Live House 依赖夜间开放、噪声、演出审批和酒水，不适合做首期主假设。",
        "adjustment": "主方案改成国潮轻餐 + 文创 + 亲子烘焙体验；Live House 降级为周末小演出、品牌日或节庆组件。",
        "operating_moves": [
            "白天做饮品、面包、文创零售和亲子 DIY，先把稳定客流吃下来。",
            "节庆日引入品牌日、小型演出或快闪，不把夜间场景常态化。",
            "与南门露天剧场联动，形成“草坪活动前后补给 + 品牌体验”闭环。",
        ],
        "price": "饮品/面包 15-45 元；亲子 DIY/文创 39-129 元；演出活动另算。",
        "evidence": ["美食/票务/亲子偏好", "消费价格带 51-150 元可解释", "空间面积足够做品牌复合"],
        "boundary": "不把夜间演艺、酒水和大型音乐消费写成首期确定收益。",
    },
    {
        "name": "南门地下预埋空间",
        "area": "5417㎡",
        "priority": "先做分区样板",
        "original": "方案一运动奥莱折扣仓；方案二大型古风/科幻实景 RPG。",
        "forecast": "面积最大，也最容易误判。运动消费和补给需求存在，但一次性满铺会把招商、消防、动线、库存和人力复杂度同时拉高。",
        "adjustment": "拆成运动户外快闪仓、赛事补给、儿童托管/休息、轻体验和招商样板区；RPG 只保留为局部节假日活动。",
        "operating_moves": [
            "先切出样板区做运动补给、装备快闪、休息/托管和赛事服务。",
            "用 30-60 天观察动线、排队、停留和品牌反馈，再决定是否扩大。",
            "把仓储、售卖和体验分开，避免地下空间一开始变成低效大卖场。",
        ],
        "price": "零售折扣按品牌定价；体验/托管 29-99 元；大型沉浸项目不进入首期预测。",
        "evidence": ["运动健身 TGI 125 / 瑜伽 TGI 382", "地下空间面积充足", "运动类 POI 供给活跃"],
        "boundary": "不按 5417㎡ 一次性满铺；不把大型 RPG 当首轮收入主线。",
    },
    {
        "name": "奥运廉洁主题展馆",
        "area": "575.72㎡",
        "priority": "第二批资质合作试点",
        "original": "身心疗愈中心，引入检测设备、疗愈室、功法/瑜伽、营养咨询和茶吧。",
        "forecast": "健康、瑜伽、康养需求有数据支撑，但医疗化和检测化会引入资质、责任和信任门槛；更适合先做健康生活方式场景。",
        "adjustment": "首期做健康展陈、八段锦/瑜伽、营养茶吧、公益讲座和企业健康活动；检测/疗程作为二期合作专区。",
        "operating_moves": [
            "雨天、高温和淡季承接室内课程与团体活动。",
            "以非医疗化表达组织内容：运动恢复、健康生活、营养茶饮、公益讲座。",
            "把企业健康日、银发课程、研学活动作为预约型产品。",
        ],
        "price": "体验/课程 49-199 元；检测/疗程只在资质合作方确认后另行设计。",
        "evidence": ["健康养生 TGI 123", "运动健身 TGI 125", "雨天/高温需要室内替代"],
        "boundary": "首期不包装医疗检测、诊疗结论和疗程销售。",
    },
    {
        "name": "10#2A03分区管理中心",
        "area": "1804.5㎡",
        "priority": "第二批文化康养试点",
        "original": "国医国学文化中心，含中医诊疗、康复理疗、国学教育、茶道、文创。",
        "forecast": "国医国学方向可成立，但诊疗、针灸、中药房属于高资质业务；首期应以文化体验和非医疗化康养课程切入。",
        "adjustment": "改成国医文化展陈、茶道/书法/八段锦、艾草香囊/文创、轻理疗体验；医疗服务只与合规机构做二期专区。",
        "operating_moves": [
            "工作日做团体课程和企业活动，周末做亲子/银发体验。",
            "文创、茶饮、香囊和课程形成低风险成交闭环。",
            "医疗相关空间单独做资质清单和责任边界，不和文化体验混写。",
        ],
        "price": "文化课程 39-199 元；文创/茶饮 20-150 元；诊疗类不进入首期预测。",
        "evidence": ["教育/亲子偏好", "健康偏好", "2A03 管理用房在南园 DWG 识别"],
        "boundary": "不把诊疗、针灸、中药房写入首期经营预测。",
    },
]


def set_run_font(run: Any, size: float = 10.5, *, bold: bool = False, color: str = INK) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(paragraph: Any, *, before: float = 0, after: float = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell: Any, top: int = 120, start: int = 140, bottom: int = 120, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell: Any, text: str, *, bold: bool = False, size: float = 9.3, color: str = INK) -> None:
    cell.text = ""
    set_cell_margin(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    set_para_spacing(p, after=1.5, line=1.08)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def set_table_widths(table: Any, widths: list[float]) -> None:
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
        for cell in table.columns[idx].cells:
            cell.width = Inches(width)


def add_title(doc: Document, basis: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=4, line=1.0)
    r = p.add_run("奥森商业改造预测与调整建议")
    set_run_font(r, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(subtitle, after=12, line=1.05)
    r = subtitle.add_run("基于区域客流、消费偏好、策划方案、图纸信息与周边供给分析")
    set_run_font(r, size=10.5, color=MUTED)

    meta = doc.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(meta, [2.1, 2.1, 2.1])
    values = [
        ("报告定位", "决策沟通版"),
        ("分析口径", "现有资料预测"),
        ("报告日期", basis["generated_at"][:10]),
    ]
    for idx, (label, value) in enumerate(values):
        cell = meta.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, f"{label}\n{value}", bold=idx == 0, size=9.2, color=BLUE if idx == 0 else INK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    style_name = f"Heading {max(1, min(level, 3))}"
    p = doc.add_paragraph(style=style_name)
    if level == 1:
        set_para_spacing(p, before=15, after=7, line=1.05)
        size, color = 16, BLUE
    elif level == 2:
        set_para_spacing(p, before=10, after=5, line=1.05)
        size, color = 12.8, BLUE_2
    else:
        set_para_spacing(p, before=6, after=3, line=1.05)
        size, color = 11.2, GREEN
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)


def add_body(doc: Document, text: str, *, size: float = 10.5, bold: bool = False, color: str = INK) -> None:
    p = doc.add_paragraph()
    set_para_spacing(p, after=6, line=1.17)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_bullet(doc: Document, text: str, *, color: str = INK, size: float = 10.1) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, after=3.5, line=1.12)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)


def add_callout(doc: Document, title: str, text: str, *, fill: str = LIGHT_GREEN, color: str = INK) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margin(cell, top=150, bottom=150, start=170, end=170)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=3, line=1.1)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=color)
    p2 = cell.add_paragraph()
    set_para_spacing(p2, after=2, line=1.15)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.2, color=color)


def add_compact_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, size=8.8, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_shading(cells[idx], WHITE)
            set_cell_text(cells[idx], value, size=8.7)


def evidence_line(basis: dict[str, Any]) -> dict[str, str]:
    h = basis["metrics"]["headline"]
    return {
        "flow": f"{h['peak_daily_all']}；{h['peak_hour_all']}",
        "coffee": h["coffee"],
        "cold_drink": h["cold_drink"],
        "yoga": h["yoga"],
        "food_visit": h["food_visit"],
        "food_price": h["food_price"],
        "poi": "、".join(f"{name} {count}" for name, count in basis["poi"]["top_categories"][:7]),
        "feature": (
            f"{basis['feature_pool']['count']} 条场景，"
            f"{basis['feature_pool']['persona_count']} 类人群，"
            f"{basis['feature_pool']['time_band_count']} 个时段，"
            f"{basis['feature_pool']['weather_count']} 类天气"
        ),
    }


def build_evidence_map(basis: dict[str, Any]) -> dict[str, Any]:
    return {
        "generated_at": basis["generated_at"],
        "purpose": "约束平台 DOCX 报告，防止无依据写作和模板化补资料主文。",
        "local_evidence_used": [
            {
                "layer": "客流与消费",
                "files": [
                    "40_quality_evidence/evidence_ledger.csv",
                    "70_outputs/processed_tables/osen_real_calibration_inputs_20260607.csv",
                ],
                "used_for": "确定首批应围绕高峰客流承接、轻餐饮、冷饮、咖啡、亲子、运动和康养，而不是直接做重资产满铺。",
            },
            {
                "layer": "策划案与节点",
                "files": ["CAD图及其计划/奥森重点项目策划思路20260521.docx"],
                "used_for": "保留用户给出的六个重点节点和原计划，但把每个节点转成预测、调整、运营动作和边界。",
            },
            {
                "layer": "CAD/图纸",
                "files": [
                    "40_quality_evidence/cad_dxf_analysis_20260605.md",
                    "40_quality_evidence/cad_pdf_proxy_analysis_20260605.md",
                ],
                "used_for": "识别白房子、南门露天剧场、廉洁馆、2A03 等图纸锚点；不把 CAD 坐标伪装成已完成 GIS 坐标。",
            },
            {
                "layer": "人物与行为特征池",
                "files": ["70_outputs/processed_tables/person_simulation_feature_derivatives_1000_20260604.csv"],
                "used_for": "把游客从标签扩展为目标、预算、时段、天气、节点和动作链，形成 state -> behavior -> demand -> advice。",
            },
            {
                "layer": "POI/供给",
                "files": ["70_outputs/processed_tables/poi_supply_candidates_amap_boundary_filter.csv"],
                "used_for": "判断周边供给活跃品类和价格带，避免只凭策划案判断业态。",
            },
        ],
        "method_basis_used": [
            {
                "source": "老板六份方法资料吸收台账",
                "files": [
                    "10_research/boss_method_materials_20260604/boss_model_inventory_20260604.md",
                    "10_research/boss_method_materials_20260604/unified_simulation_method_matrix_20260604.md",
                    "10_research/boss_method_materials_20260604/method_absorption_landing_register_20260604.md",
                ],
                "used_for": "禁用裸分数和 LLM 自由编故事，采用人群状态、行为程序、空间/消费约束和执行校核。",
            },
            {
                "source": "近年商业选址与 Agent/UX 资料护栏",
                "files": ["10_research/osen_prediction_recent_literature_guard_20260607.md"],
                "used_for": "把节点表达改为优先级、理由、动作和边界，不用单一分数或模板风险清单替代建议。",
            },
            {
                "source": "DOCX 商业报告写作格式吸收记录",
                "files": ["10_research/report_docx_writing_format_research_20260606.md"],
                "used_for": "执行摘要前置、主文服务决策、表格只放对比信息，避免机器日志和长表格。",
            },
        ],
        "external_sources_for_frame_only": [
            "GOV.UK The Green Book 2026: strategic/economic/commercial/financial/management perspectives",
            "GOV.UK business case guidance: business case as decision support, not a data dump",
            "Purdue OWL / Nottingham / Massey business report guides: executive summary, analysis, recommendations, appendices",
            "2026/2025 retail site selection and agent simulation papers recorded in local guardrail files",
        ],
        "non_claims": [
            "不宣称最终 ROI、最终收益、最终排名或完整仿真已完成。",
            "不把宏观收入直接推导为奥森节点成交能力。",
            "不让 DeepSeek 或任何 LLM 直接成为最终裁判。",
            "不把 CAD 原始图纸坐标当成已校准 GIS 坐标。",
        ],
    }


def build_markdown(basis: dict[str, Any]) -> str:
    e = evidence_line(basis)
    lines = [
        "# 奥森商业改造预测与调整建议",
        "",
        "## 执行摘要",
        "基于当前已给材料，奥森商业改造应先做低改造、可运营、可校核的轻消费与预约型场景，而不是一次性扩大面积或包装高风险业态。首批建议围绕南门露天剧场、桃花源白房子、12# 西分区管理中心推进；地下空间先做样板区，廉洁馆与 2A03 进入第二批资质合作与文化康养验证。",
        "",
        "## 核心依据",
        f"- 客流：{e['flow']}。",
        f"- 偏好：{e['coffee']}；{e['cold_drink']}；{e['yoga']}。",
        f"- 消费与竞品：{e['food_visit']}；{e['food_price']}。",
        f"- 供给与场景池：{e['poi']}；{e['feature']}。",
        "",
        "## 节点调整",
    ]
    for item in NODE_DECISIONS:
        lines.extend(
            [
                f"### {item['name']} | {item['priority']}",
                f"- 预测：{item['forecast']}",
                f"- 调整：{item['adjustment']}",
                f"- 价格带：{item['price']}",
                f"- 边界：{item['boundary']}",
            ]
        )
    lines.extend(
        [
            "",
            "## 90 天推进",
            "- 1-30 天：南门露天剧场和桃花源白房子先跑移动咖啡、冷饮、烘焙轻餐和预约活动。",
            "- 31-60 天：12# 做国潮轻餐和亲子/文创快闪，地下空间只开放样板区。",
            "- 61-90 天：廉洁馆与 2A03 进入课程、茶吧、轻康养体验验证。",
        ]
    )
    return "\n".join(lines) + "\n"


def add_executive_summary(doc: Document, basis: dict[str, Any]) -> None:
    e = evidence_line(basis)
    add_callout(
        doc,
        "本轮判断",
        "当前材料已经足以支持“先做轻消费与预约型试点”的方向：南门露天剧场、桃花源白房子、12# 西分区管理中心先推进；地下空间、医疗化康养和夜间演艺不进入首批重投入。",
        fill=LIGHT_GREEN,
        color=INK,
    )
    add_heading(doc, "1. 执行摘要", 1)
    add_body(
        doc,
        "本报告基于现有资料做预测、调整和落地排序。奥森当前最大机会是把既有高峰客流转化为可停留、可购买、可预约的轻消费场景；最大的误区是把大空间、高客流和宏观消费能力直接翻译成重资产项目。",
    )
    add_bullet(doc, "首批方向：咖啡/冷饮、烘焙轻餐、草坪活动、亲子小活动、国潮轻餐与文创快闪。")
    add_bullet(doc, "首批节点：南门露天剧场、桃花源白房子、12# 西分区管理中心。")
    add_bullet(doc, "审慎推进：地下预埋空间只做分区样板；廉洁馆与 2A03 先做非医疗化文化康养。")
    add_bullet(doc, "表达方式：不用裸分数替代判断；用“为什么成立、怎样先做、什么不要先做”服务业务决策。")

    doc.add_page_break()
    add_heading(doc, "2. 已有材料支撑的判断", 1)
    add_compact_table(
        doc,
        ["材料层", "关键线索", "对方案的约束"],
        [
            ["客流", e["flow"], "高峰流量足够，但报告只把它作为承接机会，不直接推导收益。"],
            ["偏好", f"{e['coffee']}；{e['cold_drink']}；{e['yoga']}", "咖啡、冷饮、运动/康养可进入候选，但必须按节点、时段和天气组织。"],
            ["消费", e["food_price"], "价格带宜分为日常轻消费、活动消费、预约课程；不宜只押高客单。"],
            ["供给", e["poi"], "周边供给活跃，园内应做差异化停留和场景组合，而不是复制门店。"],
            ["行为池", e["feature"], "建议要覆盖人群、预算、时段、天气和动作链，不能只写“亲子/白领/游客”。"],
        ],
        [1.05, 2.65, 2.6],
    )


def add_strategy(doc: Document) -> None:
    add_heading(doc, "3. 推荐商业组合", 1)
    add_body(
        doc,
        "六个节点不应各自孤立招商。更稳妥的做法是把它们合成三组商业动作：先让低改造、强场景、能试运营的节点跑起来，再用真实反馈决定是否扩大。",
    )
    add_compact_table(
        doc,
        ["组合", "包含节点", "主要客群与动作", "先看什么结果"],
        [
            [
                "南门轻消费与活动簇",
                "南门露天剧场、12#、地下空间样板区",
                "亲子家庭、游客、运动人群；咖啡冷饮、烘焙、草坪活动、快闪仓、赛事补给",
                "时段客流、购买转化、排队、停留、天气影响和活动复购",
            ],
            [
                "北园景观打卡簇",
                "桃花源白房子",
                "花海拍照、生日/团建、轻餐下午茶、摄影预约",
                "预约率、包场成交、花季波峰、平日空置率",
            ],
            [
                "室内文化康养簇",
                "廉洁馆、10#2A03",
                "银发、白领、研学和亲子；八段锦、茶吧、书法、健康讲座、文创",
                "到课率、复购、团体活动、资质合作可行性",
            ],
        ],
        [1.35, 1.45, 2.05, 1.75],
    )


def add_node_cards(doc: Document) -> None:
    add_heading(doc, "4. 六个节点的预测与调整", 1)
    add_body(
        doc,
        "以下不是最终排名，而是基于当前资料生成的推进优先级。每个节点都按“原计划理解 -> 预测判断 -> 调整动作 -> 价格/时段 -> 边界”落地，便于后续用真实试运营数据继续修正。",
        color=MUTED,
    )
    for idx, item in enumerate(NODE_DECISIONS, 1):
        add_heading(doc, f"4.{idx} {item['name']} | {item['priority']}", 2)
        add_callout(
            doc,
            "预测判断",
            item["forecast"],
            fill=LIGHT_BLUE if "第一批" in item["priority"] else LIGHT_GOLD,
            color=INK,
        )
        add_body(doc, f"原计划理解：{item['original']}", size=10.2, color=MUTED)
        add_body(doc, f"调整建议：{item['adjustment']}", bold=True)
        for move in item["operating_moves"]:
            add_bullet(doc, move)
        add_body(doc, f"价格与时段口径：{item['price']}", color=GREEN)
        add_body(doc, "支撑线索：" + "；".join(item["evidence"]) + "。", size=10.0, color=MUTED)
        add_body(doc, f"边界：{item['boundary']}", size=10.0, color=GOLD)


def add_behavior_forecast(doc: Document, basis: dict[str, Any]) -> None:
    examples = basis["feature_pool"].get("examples", [])[:6]
    add_heading(doc, "5. 人群行为预测如何进入建议", 1)
    add_body(
        doc,
        "本轮不把人群写成空泛标签，而是采用“状态 -> 行为 -> 需求 -> 节点动作”的链条。这样能解释为什么同样是高客流，有些空间适合快取，有些适合预约，有些只能先做样板。",
    )
    rows = []
    for row in examples:
        rows.append(
            [
                row.get("persona_name", ""),
                row.get("time_band_name", "") + " / " + row.get("weather_name", ""),
                row.get("node_context_name", ""),
                row.get("action", ""),
            ]
        )
    add_compact_table(
        doc,
        ["人群状态", "时段/天气", "空间语境", "更合理的动作"],
        rows[:6],
        [1.45, 1.25, 1.25, 2.35],
    )
    add_body(
        doc,
        "对报告的直接影响：晨练人群更适合补水、快取和路线不中断；亲子家庭更在意安全、座位、遮阴和轻食；拍照/游客更容易被景观、花季和可分享产品触发；银发和白领对康养场景更看重信任、低风险和预约秩序。",
    )


def add_implementation(doc: Document) -> None:
    add_heading(doc, "6. 90 天推进节奏", 1)
    add_compact_table(
        doc,
        ["阶段", "主动作", "判断口径"],
        [
            [
                "1-30 天",
                "南门露天剧场和桃花源白房子先跑移动咖啡、冷饮、烘焙轻餐、花季/周末预约活动。",
                "购买转化、停留时长、排队、天气影响、客诉、复购线索。",
            ],
            [
                "31-60 天",
                "12# 做国潮轻餐与亲子/文创快闪；地下空间只开放运动补给和快闪样板区。",
                "品牌吸引力、亲子停留、样板区动线、地下空间运营压力。",
            ],
            [
                "61-90 天",
                "廉洁馆与 2A03 做课程、茶吧、轻康养体验；沉淀第二批招商条件。",
                "预约率、到课率、课程复购、银发/白领接受度、资质合作边界。",
            ],
        ],
        [1.05, 3.35, 1.95],
    )

def add_appendix(doc: Document, basis: dict[str, Any]) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附录 A. 数据口径与实施边界", 1)
    add_body(
        doc,
        "以下说明用于帮助后续讨论统一口径。它不改变本报告的主线判断，也不替代现场踏勘、深化设计、招商测算和运营试点。",
    )
    add_body(
        doc,
        "执行上，建议先用小样本试运营校准价格带、排队、停留、复购和天气敏感性；不把宏观收入、方案测算假设或周边客单直接写成奥森节点收益。",
        size=9.8,
        color=MUTED,
    )
    add_compact_table(
        doc,
        ["口径", "本报告如何使用", "实施前需要控制的边界"],
        [
            [
                "客流与消费偏好",
                "用于判断哪些业态和时段更有承接机会。",
                "客流和偏好不等于成交率，收益仍需试运营流水校准。",
            ],
            [
                "策划方案与图纸",
                "用于识别重点节点、空间角色和可先行测试的商业动作。",
                "图纸信息进入施工和招商前，应由设计、消防、运营团队复核。",
            ],
            [
                "周边供给与价格带",
                "用于约束园内业态组合和初始价格区间。",
                "周边客单不能直接等同于园内成交能力。",
            ],
            [
                "节点优先级",
                "用于安排首批试点、第二批验证和暂缓扩大项。",
                "优先级不是最终投资批复，仍需结合审批、招商、工程和运营排期。",
            ],
        ],
        [1.6, 2.4, 2.2],
    )
    add_heading(doc, "附录 B. 本轮建议的使用方式", 1)
    add_body(
        doc,
        "建议把本报告作为方案确认和试运营设计的第一版工作稿：先确认首批试点方向，再组织现场复核、招商沟通和 30-90 天试运营。试运营结果返回后，可继续修正价格带、排班、节点组合和招商边界。",
    )
    for src in [
        "方案确认：确认南门、桃花源白房子、12# 是否作为首批试点组合。",
        "现场复核：核对动线、座位、遮阴、排队、供电、排水、垃圾回收和安全边界。",
        "招商沟通：优先接触咖啡/冷饮、烘焙轻餐、国潮轻餐、亲子文创、运动补给和轻康养合作方。",
        "试运营设计：按 30 天一个周期记录客流、购买、预约、排队、天气影响和客诉。",
        "复盘调整：用试运营结果决定地下空间、医疗化康养和夜间活动是否扩大。",
    ]:
        add_bullet(doc, src, size=9.5, color=MUTED)
    add_body(doc, f"报告日期：{basis['generated_at'][:10]}", size=9.4, color=MUTED)


def write_docx(basis: dict[str, Any], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = FONT_LATIN
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.color.rgb = RGBColor.from_string(BLUE)

    add_title(doc, basis)
    add_executive_summary(doc, basis)
    add_strategy(doc)
    add_node_cards(doc)
    add_behavior_forecast(doc, basis)
    add_implementation(doc)
    add_appendix(doc, basis)
    doc.save(path)


def build_outputs(*, out_dir: Path | None = None, copy_static: bool = True) -> dict[str, Any]:
    target_dir = out_dir or OUT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    QUALITY_DIR.mkdir(parents=True, exist_ok=True)
    if copy_static:
        STATIC_DIR.mkdir(parents=True, exist_ok=True)

    docx_path = target_dir / DOCX_OUT.name
    md_path = target_dir / MD_OUT.name
    basis_path = BASIS_OUT
    evidence_map_path = EVIDENCE_MAP_OUT

    basis = build_basis()
    basis["generated_at"] = datetime.now().isoformat(timespec="seconds")
    basis["report_redesign_basis"] = {
        "design_preset": "standard_business_brief",
        "platform_route": "/api/reports/site-selection/download?format=docx",
        "main_principle": "用户已给材料优先；外部资料约束写法和方法边界，不替代本地证据。",
        "no_claims": [
            "不宣称最终 ROI",
            "不宣称完整仿真完成",
            "不把宏观收入直接推导为节点收益",
            "不把 CAD 坐标当作已校准 GIS 坐标",
        ],
    }
    evidence_map = build_evidence_map(basis)

    write_docx(basis, docx_path)
    md_path.write_text(build_markdown(basis), encoding="utf-8")
    basis_path.write_text(json.dumps(basis, ensure_ascii=False, indent=2), encoding="utf-8")
    evidence_map_path.write_text(json.dumps(evidence_map, ensure_ascii=False, indent=2), encoding="utf-8")

    if copy_static:
        shutil.copy2(docx_path, STATIC_DOCX)
        shutil.copy2(basis_path, STATIC_BASIS)

    return {
        "status": "pass",
        "output_status": "needs_review",
        "docx": str(docx_path.resolve()),
        "docx_bytes": docx_path.stat().st_size,
        "markdown": str(md_path.resolve()),
        "basis": str(basis_path.resolve()),
        "evidence_map": str(evidence_map_path.resolve()),
        "download_filename": DOCX_OUT.name,
        "platform_route": "/api/reports/site-selection/download?format=docx",
        "static_copy": str(STATIC_DOCX.resolve()) if copy_static else "",
    }


def main() -> int:
    result = build_outputs()
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
