from __future__ import annotations

import csv
import html
import json
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "80_delivery"
QUALITY_DIR = ROOT / "40_quality_evidence"
STATIC_DIR = ROOT / "90_p6_expert_dashboard" / "static"

EVIDENCE_LEDGER = QUALITY_DIR / "evidence_ledger.csv"
CALIBRATION_CSV = ROOT / "70_outputs" / "processed_tables" / "osen_real_calibration_inputs_20260607.csv"
FEATURE_CSV = ROOT / "70_outputs" / "processed_tables" / "person_simulation_feature_derivatives_1000_20260604.csv"
POI_CSV = ROOT / "70_outputs" / "processed_tables" / "poi_supply_candidates_amap_boundary_filter.csv"
PLAN_DOCX = ROOT / "CAD图及其计划" / "奥森重点项目策划思路20260521.docx"
CAD_DXF_JSON = QUALITY_DIR / "cad_dxf_analysis_20260605.json"
CAD_PDF_JSON = QUALITY_DIR / "cad_pdf_proxy_analysis_20260605.json"
METHOD_BASIS = QUALITY_DIR / "osen_report_repair_evidence_basis_20260607.md"


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return [dict(row) for row in csv.DictReader(f)]


def read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def docx_text(path: Path) -> str:
    if not path.exists():
        return ""
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text = re.sub(r"<[^>]+>", "", xml)
    return text.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")


def evidence_by_id(rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    return {str(row.get("metric_id") or ""): row for row in rows}


def ev_text(row: dict[str, str]) -> str:
    return f"{row.get('metric_name', '')}：{row.get('value', '')}{row.get('unit', '')}"


def metric_pack(rows: list[dict[str, str]]) -> dict[str, Any]:
    by_id = evidence_by_id(rows)
    pick_ids = [
        "MET-0004",
        "MET-0005",
        "MET-0006",
        "MET-0007",
        "MET-0008",
        "MET-0019",
        "MET-0020",
        "MET-0021",
        "MET-0022",
        "MET-0025",
        "MET-0026",
        "MET-0027",
        "MET-0028",
        "MET-0049",
        "MET-0050",
        "MET-0257",
        "MET-0258",
        "MET-0259",
        "MET-0260",
    ]
    return {
        "selected_evidence": [
            {
                "metric_id": mid,
                "metric_name": by_id.get(mid, {}).get("metric_name", ""),
                "value": by_id.get(mid, {}).get("value", ""),
                "unit": by_id.get(mid, {}).get("unit", ""),
                "source_file": by_id.get(mid, {}).get("source_file", ""),
            }
            for mid in pick_ids
            if mid in by_id
        ],
        "headline": {
            "peak_daily_all": ev_text(by_id.get("MET-0004", {})),
            "peak_hour_all": ev_text(by_id.get("MET-0007", {})),
            "coffee": f"{ev_text(by_id.get('MET-0025', {}))}；{ev_text(by_id.get('MET-0026', {}))}",
            "cold_drink": f"{ev_text(by_id.get('MET-0027', {}))}；{ev_text(by_id.get('MET-0028', {}))}",
            "yoga": f"{ev_text(by_id.get('MET-0049', {}))}；{ev_text(by_id.get('MET-0050', {}))}",
            "food_price": "；".join(
                ev_text(by_id.get(mid, {}))
                for mid in ["MET-0020", "MET-0258", "MET-0260"]
                if by_id.get(mid)
            ),
            "food_visit": "；".join(
                ev_text(by_id.get(mid, {}))
                for mid in ["MET-0019", "MET-0257", "MET-0259"]
                if by_id.get(mid)
            ),
        },
    }


def calibration_pack(rows: list[dict[str, str]]) -> dict[str, Any]:
    wanted = ["ORCI-001", "ORCI-002", "ORCI-003", "ORCI-101", "ORCI-102", "ORCI-103", "ORCI-106", "ORCI-107"]
    by_id = {row.get("calibration_id", ""): row for row in rows}
    return {
        "count": len(rows),
        "items": [
            {
                "calibration_id": cid,
                "indicator_name": by_id.get(cid, {}).get("indicator_name", ""),
                "value": by_id.get(cid, {}).get("value", ""),
                "unit": by_id.get(cid, {}).get("unit", ""),
                "simulation_use": by_id.get(cid, {}).get("simulation_use", ""),
            }
            for cid in wanted
            if cid in by_id
        ],
    }


def feature_pack(rows: list[dict[str, str]]) -> dict[str, Any]:
    persona_order = ["runner", "parent_family", "office_lunch", "tourist_photo", "culture_visitor", "silver_senior", "cyclist", "night_event"]
    examples = []
    for pid in persona_order:
        row = next((item for item in rows if item.get("persona_id") == pid), None)
        if row:
            examples.append(
                {
                    "persona_id": pid,
                    "persona_name": row.get("persona_name", ""),
                    "state": row.get("persona_core_need", ""),
                    "income_price_band": row.get("income_price_band", ""),
                    "time_band_name": row.get("time_band_name", ""),
                    "weather_name": row.get("weather_name", ""),
                    "node_context_name": row.get("node_context_name", ""),
                    "action": row.get("candidate_supply_action_name", ""),
                }
            )
    return {
        "count": len(rows),
        "persona_count": len({row.get("persona_id") for row in rows}),
        "income_segment_count": len({row.get("income_segment_id") for row in rows}),
        "time_band_count": len({row.get("time_band_id") for row in rows}),
        "weather_count": len({row.get("weather_id") for row in rows}),
        "examples": examples,
    }


def poi_pack(rows: list[dict[str, str]]) -> dict[str, Any]:
    osen = [row for row in rows if row.get("park_id") == "sample_olympic_forest"]
    categories: dict[str, int] = {}
    for row in osen:
        categories[row.get("standard_categories", "")] = categories.get(row.get("standard_categories", ""), 0) + 1
    return {
        "count": len(osen),
        "top_categories": sorted(categories.items(), key=lambda item: item[1], reverse=True)[:12],
        "priced_examples": [
            {
                "poi_name": row.get("poi_name", ""),
                "category": row.get("standard_categories", ""),
                "cost_yuan": row.get("cost_yuan", ""),
                "distance_band": row.get("distance_band", ""),
                "boundary_filter_status": row.get("boundary_filter_status", ""),
            }
            for row in osen
            if row.get("cost_yuan")
        ][:10],
    }


NODES = [
    {
        "node": "桃花源白房子",
        "area": "155㎡",
        "original": "活动+轻餐，多功能活动据点，承接婚礼/摄影、亲子/团建、咖啡茶饮和季节下午茶。",
        "prediction": "需求成立，但应以轻餐、摄影预约、生日/团建小单和花季活动为主；高固定投入和重厨房不宜先上。",
        "adjustment": "保留“白房子+花海+轻餐”定位，取消一开始重改阳光房的冲动，先做可移动外摆、预约活动、咖啡/冷饮和周末亲子小活动。",
        "price": "日常 20-60 元；预约活动/冷餐会按 80-180 元/人做情景，不写成确定客单。",
        "schedule": "春秋花期和周末为主；平日以咖啡、轻食、办公/活动公司后台支撑为辅。",
        "priority": "第一批低改造试点",
        "method": "HumanLM: 花海打卡/亲子/活动客的目标与情绪；ROTE: 到达-拍照-休息-轻餐-离开/预约；SSR: 由理由映射为优先推进。",
    },
    {
        "node": "奥运廉洁主题展馆",
        "area": "575.72㎡",
        "original": "身心疗愈中心，引入检测设备、疗愈室、功法/瑜伽、营养咨询和茶吧。",
        "prediction": "健康、瑜伽和康养需求有数据支持，但医疗化、检测化、疗程包风险高；更适合先做健康生活方式和轻康养体验。",
        "adjustment": "把“德国精密检测+高端疗程”降级为二期，首期做健康展陈、八段锦/瑜伽、营养茶吧、公益讲座和企业健康活动。",
        "price": "体验/课程 49-199 元；检测/疗程只在资质合作方确认后另行设计。",
        "schedule": "全天候预约制，雨雪高温时作为室内替代空间。",
        "priority": "第二批资质合作试点",
        "method": "HumanLM: 银发/白领的健康焦虑与信任门槛；ROTE: 参观-咨询-体验-预约；双重奖励: 微观行为合理与宏观健康偏好一致。",
    },
    {
        "node": "12#西分区管理中心",
        "area": "1059.39㎡",
        "original": "方案一 Live House；方案二北冰洋/义利国潮轻餐、文创和亲子体验。",
        "prediction": "北冰洋/义利方案更稳；Live House 依赖夜间开放、噪声、演出审批和酒水，不适合做首期主假设。",
        "adjustment": "主方案改成白天国潮轻餐+文创+亲子烘焙体验；Live House 降级为周末小演出、品牌日或节庆活动组件。",
        "price": "饮品/面包 15-45 元；亲子 DIY/文创 39-129 元；演出活动另算。",
        "schedule": "白天和周末为主；夜间仅在审批成立时做活动日。",
        "priority": "第一批品牌联名试点",
        "method": "SSR: 两方案先写理由再映射优先级；Logit/Huff: 白天稳定客群优于夜间不确定客群。",
    },
    {
        "node": "南门地下预埋空间",
        "area": "5417㎡",
        "original": "方案一运动奥莱折扣仓；方案二大型古风/科幻实景 RPG。",
        "prediction": "面积大但重资产风险最高；运动消费线索存在，但一次性做大型奥莱或 RPG 都容易被客流、消防、招商和运营复杂度拖垮。",
        "adjustment": "首期不做满铺。拆成运动户外快闪仓、赛事补给、儿童托管/休息、轻体验和招商样板区；RPG 只保留为节假日预约活动或局部场景。",
        "price": "零售折扣按品牌定价；体验/托管 29-99 元；大型沉浸式项目不进入首期收入预测。",
        "schedule": "节假日、赛事日和暑期为主，平日做仓储/品牌快闪/课程预约。",
        "priority": "先做分区样板，不进入一批满铺",
        "method": "ROTE: 运动家庭-购物-儿童托管-离开；RL+LLM: 动作空间复杂，先用规则和小样本反馈约束。",
    },
    {
        "node": "南门露天剧场",
        "area": "草坪/湖畔空间",
        "original": "湖畔阳光烘焙工坊 + 四季草坪生活场，含咖啡烘焙、市集、音乐、亲子工坊。",
        "prediction": "与奥森数据最匹配：咖啡、冷饮、美食、亲子和草坪停留都能解释；适合作为第一批标杆试点。",
        "adjustment": "做“烘焙+咖啡+草坪活动+移动补给”的组合，不先做高噪声演出；活动日历和天气预案要一起设计。",
        "price": "咖啡/冷饮 20-45 元；烘焙轻餐 35-98 元；草坪套餐/亲子活动 68-168 元。",
        "schedule": "周末、节假日、下午和活动前后最强；高温/雨天切换为室内烘焙和轻活动。",
        "priority": "第一批主推试点",
        "method": "HumanLM: 亲子/拍照/游客的停留状态；ROTE: 湖边停留-闻到烘焙-购买-草坪消费；宏观校准: 咖啡/冷饮 TGI 与峰值客流。",
    },
    {
        "node": "10#2A03分区管理中心",
        "area": "1804.5㎡",
        "original": "国医国学文化中心，含中医诊疗、康复理疗、国学教育、茶道、文创。",
        "prediction": "国医国学方向可成立，但诊疗、针灸、中药房属于高资质业务；首期应以文化体验和非医疗化康养课程切入。",
        "adjustment": "改成国医文化展陈、茶道/书法/八段锦、艾草香囊/文创、轻理疗体验；医疗服务只与合规机构做二期专区。",
        "price": "文化课程 39-199 元；文创/茶饮 20-150 元；诊疗类不进入首期预测。",
        "schedule": "工作日可做团体课程，周末做亲子/银发体验。",
        "priority": "第二批文化康养试点",
        "method": "HumanLM: 银发/文化客的信任与低风险需求；ROTE: 参观-体验-购买/预约；许可边界作为动作 mask。",
    },
]


def build_basis() -> dict[str, Any]:
    evidence_rows = read_csv(EVIDENCE_LEDGER)
    calibration_rows = read_csv(CALIBRATION_CSV)
    feature_rows = read_csv(FEATURE_CSV)
    poi_rows = read_csv(POI_CSV)
    cad_dxf = read_json(CAD_DXF_JSON)
    cad_pdf = read_json(CAD_PDF_JSON)
    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "status": "needs_review",
        "source_files": {
            "evidence_ledger": str(EVIDENCE_LEDGER.relative_to(ROOT)),
            "calibration_csv": str(CALIBRATION_CSV.relative_to(ROOT)),
            "feature_derivatives": str(FEATURE_CSV.relative_to(ROOT)),
            "poi_candidates": str(POI_CSV.relative_to(ROOT)),
            "plan_docx": str(PLAN_DOCX.relative_to(ROOT)),
            "cad_dxf_analysis": str(CAD_DXF_JSON.relative_to(ROOT)),
            "cad_pdf_proxy": str(CAD_PDF_JSON.relative_to(ROOT)),
            "method_basis": str(METHOD_BASIS.relative_to(ROOT)),
        },
        "metrics": metric_pack(evidence_rows),
        "calibration": calibration_pack(calibration_rows),
        "feature_pool": feature_pack(feature_rows),
        "poi": poi_pack(poi_rows),
        "plan_text_chars": len(docx_text(PLAN_DOCX)),
        "cad": {
            "dxf_result_count": len(cad_dxf.get("results", [])),
            "pdf_hits": cad_pdf.get("hits", []),
            "cad_boundary": cad_pdf.get("use_boundary", ""),
        },
        "nodes": NODES,
        "method_landing": [
            "HumanLM: 每个节点先判断人群潜在状态，而不是只写客群标签。",
            "ROTE: 每个节点给出触发-行为-消费-放弃的行为程序。",
            "SSR/FLR: 不直接给裸分，先写理由，再映射推进优先级。",
            "RL+LLM 双重奖励: 当前不训练 PPO，但报告用微观行为合理性和宏观数据一致性做校验。",
            "Agent Bank: 人物弱记忆来自 PDF/TGI/POI/策划资料，不凭空捏人。",
        ],
    }


def report_markdown(basis: dict[str, Any]) -> str:
    h = basis["metrics"]["headline"]
    feature = basis["feature_pool"]
    poi = basis["poi"]
    lines = [
        "# 奥森商业改造预测与调整方案报告",
        "",
        f"生成时间：{basis['generated_at']}",
        "",
        "## 1. 结论先行",
        "",
            "这版报告直接使用本文件夹内已给数据、策划案、图纸和老板方法，形成当前可执行的预测与调整方案。需要复核的内容放到后文风险和试运营指标中，不作为主文结论。",
        "",
        "- 第一批应优先做：南门露天剧场、桃花源白房子、12#西分区管理中心的低改造/品牌联名试点。",
        "- 应降级或分期：南门地下预埋空间、奥运廉洁主题展馆、10#2A03 分区管理中心的重医疗、重演出、重沉浸式项目。",
        "- 最强数据线索来自：峰值到访、咖啡/冷饮/美食 TGI、瑜伽/康养偏好、热门 POI 客单和策划节点面积。",
        "- 当前建议不是最终 ROI，而是“用已有数据能推出来的经营动作、价格带、时段和试运营路径”。",
        "",
        "## 2. 已给数据如何支撑预测",
        "",
        f"- 客流底座：{h['peak_daily_all']}；{h['peak_hour_all']}。这说明奥森不是小流量点，关键是把人流转成可停留、可购买、可预约的节点行为。",
        f"- 咖啡线索：{h['coffee']}。咖啡不应只作为“好看业态”，而应进入白房子、露天剧场、12# 的快取和停留场景。",
        f"- 冷饮线索：{h['cold_drink']}。夏季、运动后、亲子活动和湖边草坪场景应有冷饮/补水/冰品组合。",
        f"- 瑜伽康养线索：{h['yoga']}。康养可以做，但应先做低医疗风险的课程、体验和活动。",
        f"- 美食热门与客单线索：{h['food_visit']}；{h['food_price']}。价格带不能只看高客单，园内应按 20-60、60-120、120+ 三层设计。",
        f"- 本地 POI 供给：当前奥森相关候选 {poi['count']} 条，前列类别包括 " + "、".join(f"{name} {count}" for name, count in poi["top_categories"][:8]) + "。",
        f"- 人物仿真底座：{feature['count']} 条衍生场景，覆盖 {feature['persona_count']} 类人群、{feature['income_segment_count']} 类收入/价格带、{feature['time_band_count']} 个时段、{feature['weather_count']} 类天气。",
        "",
        "## 3. 人群行为预测",
        "",
        "这一节不再把人群压成一张密表。当前预测先看“人为什么来、在什么时段停留、愿意为什么付费、遇到什么会放弃”，再把动作映射到节点。这样更接近 HumanLM 与 ROTE 的用法，也更适合给业务方阅读。",
        "",
    ]
    for item in feature["examples"]:
        implication = {
            "runner": "入口、闸口和环线路径需要补水/低价补给，适合饮水机和移动售卖。",
            "parent_family": "草坪、湖边和可坐可遮阴空间适合烘焙、轻餐、亲子 DIY。",
            "office_lunch": "快取咖啡和轻食必须控制排队和绕行，适合 12# 或入口附近。",
            "tourist_photo": "桃花源和湖边节点应强化拍照后消费与纪念品。",
            "culture_visitor": "展馆类节点要有讲解、茶吧、文创和活动预约。",
            "silver_senior": "康养/国医节点要弱医疗化，强调舒适、可信和低风险。",
            "cyclist": "运动补给和装备快闪更适合地下空间或入口节点。",
            "night_event": "夜间活动只能作为审批成立后的活动组件。",
        }.get(item["persona_id"], "作为场景补充。")
        lines.extend(
            [
                f"### {item['persona_name']}",
                "",
                f"- 状态与动机：{item['state']}",
                f"- 价格/预算：{item['income_price_band']}",
                f"- 当前动作线索：{item['action']}；高频时段为{item['time_band_name']}，天气场景为{item['weather_name']}。",
                f"- 对节点的含义：{implication}",
                "",
            ]
        )
    lines.extend(["", "## 4. 六节点预测与调整", ""])
    for idx, node in enumerate(NODES, start=1):
        lines.extend(
            [
                f"### {idx}. {node['node']}（{node['area']}）",
                "",
                f"- 原计划：{node['original']}",
                f"- 当前预测：{node['prediction']}",
                f"- 调整方案：{node['adjustment']}",
                f"- 价格带：{node['price']}",
                f"- 运营时段：{node['schedule']}",
                f"- 推进优先级：{node['priority']}",
                f"- 方法落地：{node['method']}",
                "",
            ]
        )
    lines.extend(
        [
            "## 5. 组合推进策略",
            "",
            "| 组合 | 包含节点 | 预测作用 | 先做什么 |",
            "|---|---|---|---|",
            "| 南门轻消费与活动簇 | 南门露天剧场、南门地下预埋部分空间、12#联动 | 承接咖啡/冷饮/轻餐/草坪活动/亲子需求，是第一批最容易看见效果的组合 | 烘焙咖啡、移动补给、草坪活动、品牌快闪 |",
            "| 北园景观打卡簇 | 桃花源白房子 | 承接花海、拍照、生日/团建、轻餐停留 | 移动外摆、预约活动、低改造摄影/生日套餐 |",
            "| 室内康养文化簇 | 廉洁馆、2A03 | 承接雨天、高温、银发、白领、文化研学与健康课程 | 茶吧、讲座、八段锦/瑜伽、国医文化体验 |",
            "| 重资产观察簇 | 南门地下预埋、Live House、医疗检测 | 有想象力，但复杂度高，应作为二期而非首期主方案 | 分区样板、招商测试、审批路径确认 |",
            "",
            "## 6. 试运营设计",
            "",
            "- 第 1 阶段（2-4 周）：南门露天剧场和桃花源做咖啡/冷饮/烘焙/轻餐试点，记录时段客流、购买转化、排队、天气影响和客单区间。",
            "- 第 2 阶段（4-8 周）：12# 做北冰洋/义利联名快闪，测试亲子、怀旧和白天客群；Live House 只做小活动，不做主招商假设。",
            "- 第 3 阶段（8-12 周）：廉洁馆/2A03 做健康文化课程和茶吧，观察预约率、复购意愿和信任门槛。",
            "- 地下空间只做分区招商样板，不以 5417㎡ 满铺作为首轮预测前提。",
            "",
            "## 7. 情景预测",
            "",
            "| 情景 | 判断 | 预期结果 | 风险 |",
            "|---|---|---|---|",
            "| 保守情景 | 只做低改造咖啡/冷饮/轻餐/课程 | 能快速验证需求，投资压力低，容易调整 | 收入天花板较低 |",
            "| 推荐情景 | 南门簇 + 桃花源 + 12# 国潮轻餐同步试点 | 覆盖亲子、游客、运动、白领和活动人群，最符合当前数据 | 需要统一运营节奏和活动日历 |",
            "| 激进情景 | 地下空间、Live House、医疗检测同步重投 | 想象空间大，但对审批、招商、夜间、资质和客流转化依赖过强 | 容易形成高投入、低验证、难回收 |",
            "",
            "## 8. 风险与复核项",
            "",
            "这些不是主结论，而是执行前的复核清单：",
            "",
            "- 价格带必须用试运营支付数据校准，不能只用北京市宏观收入判断。",
            "- 地下空间、医疗检测、Live House、夜间活动需要许可、消防、安全和责任边界。",
            "- CAD 坐标需做 GIS 控制点校准，避免把工程坐标直接当地图路线。",
            "- 天气、淡旺季和活动日历必须进入运营排班，不然冷饮、烘焙、草坪活动会波动很大。",
            "",
            "## 9. 本稿依据",
            "",
            f"- 依据链文件：`{METHOD_BASIS.relative_to(ROOT)}`",
            "- 数据来源：奥森区域大数据 PDF、证据台账、真实校准输入、人物仿真衍生特征池、POI 候选、策划 DOCX、CAD/DWG/PDF 图纸。",
            "- 方法来源：HumanLM、ROTE、RL+LLM 双重奖励、SSR/FLR、Agent Bank、Huff/Logit/Gravity 的本地工程化约束。",
        ]
    )
    return "\n".join(lines) + "\n"


def add_run(paragraph: Any, text: str, *, bold: bool = False, size: int = 10, color: str = "111827") -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # type: ignore[name-defined]


def set_font(run: Any, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)  # type: ignore[name-defined]


def table(doc: Document, rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.style = "Table Grid"
    for i, cell in enumerate(t.rows[0].cells):
        cell.text = rows[0][i]
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r)
                r.bold = True
                r.font.size = Pt(8.5)
    for row in rows[1:]:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_font(r)
                    r.font.size = Pt(8)


def write_docx(basis: dict[str, Any], path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("奥森商业改造预测与调整方案报告")
    set_font(run)
    run.bold = True
    run.font.size = Pt(22)
    subtitle = doc.add_paragraph("基于已给 PDF 数据、策划 DOCX、CAD 图纸、人物仿真特征池与老板方法模型")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("4B5563")
    doc.add_paragraph(f"生成时间：{basis['generated_at']}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    md = report_markdown(basis)
    blocks = md.splitlines()
    table_buffer: list[list[str]] = []
    in_table = False
    for line in blocks[4:]:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            if table_buffer:
                table(doc, table_buffer)
                table_buffer = []
                in_table = False
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            if table_buffer:
                table(doc, table_buffer)
                table_buffer = []
                in_table = False
            doc.add_heading(line[4:], level=2)
        elif line.startswith("|") and line.endswith("|"):
            parts = [part.strip() for part in line.strip("|").split("|")]
            if all(set(part) <= {"-", " "} for part in parts):
                in_table = True
                continue
            table_buffer.append(parts)
            in_table = True
        else:
            if in_table and table_buffer:
                table(doc, table_buffer)
                table_buffer = []
                in_table = False
            if not line.strip():
                continue
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line[2:])
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
            set_font(run)
            run.font.size = Pt(10)
    if table_buffer:
        table(doc, table_buffer)
    doc.save(path)


def markdown_to_html(md: str) -> str:
    lines = []
    in_list = False
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if not table_rows:
            return
        header, *body = table_rows
        lines.append("<table><thead><tr>" + "".join(f"<th>{html.escape(c)}</th>" for c in header) + "</tr></thead><tbody>")
        for row in body:
            lines.append("<tr>" + "".join(f"<td>{html.escape(c)}</td>" for c in row) + "</tr>")
        lines.append("</tbody></table>")
        table_rows = []
        in_table = False

    for raw in md.splitlines():
        line = raw.strip()
        if in_table and not line.startswith("|"):
            flush_table()
        if in_list and not line.startswith("- "):
            lines.append("</ul>")
            in_list = False
        if not line:
            continue
        if line.startswith("|") and line.endswith("|"):
            parts = [part.strip() for part in line.strip("|").split("|")]
            if all(set(part) <= {"-", " "} for part in parts):
                continue
            table_rows.append(parts)
            in_table = True
        elif line.startswith("# "):
            lines.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith("## "):
            lines.append(f"<h2>{html.escape(line[3:])}</h2>")
        elif line.startswith("### "):
            lines.append(f"<h3>{html.escape(line[4:])}</h3>")
        elif line.startswith("- "):
            if not in_list:
                lines.append("<ul>")
                in_list = True
            lines.append(f"<li>{html.escape(line[2:])}</li>")
        else:
            lines.append(f"<p>{html.escape(line)}</p>")
    if in_list:
        lines.append("</ul>")
    flush_table()
    return "\n".join(lines)


def write_html(md: str, path: Path) -> None:
    body = markdown_to_html(md)
    page = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>奥森商业改造预测与调整方案报告</title>
  <style>
    body {{ margin: 0; background: #f6f8f7; color: #13252d; font-family: "Microsoft YaHei", Arial, sans-serif; }}
    .topbar {{ position: sticky; top: 0; z-index: 10; background: rgba(246, 248, 247, 0.96); border-bottom: 1px solid #d9e3df; backdrop-filter: blur(10px); }}
    .topbar-inner {{ max-width: 1160px; margin: 0 auto; padding: 14px 28px; display: flex; align-items: center; justify-content: space-between; gap: 14px; }}
    .topbar-title {{ font-weight: 700; color: #0f2f28; }}
    .actions {{ display: flex; gap: 10px; flex-wrap: wrap; }}
    .actions a {{ display: inline-flex; align-items: center; justify-content: center; min-height: 36px; padding: 0 14px; border-radius: 8px; border: 1px solid #b9d4ca; background: #fff; color: #0e5a46; text-decoration: none; font-weight: 650; }}
    .actions a.primary {{ background: #0e5a46; color: #fff; border-color: #0e5a46; }}
    main {{ max-width: 1160px; margin: 0 auto; padding: 36px 28px 72px; }}
    h1 {{ font-size: 34px; margin: 0 0 12px; }}
    h2 {{ margin-top: 34px; padding-top: 18px; border-top: 1px solid #d9e3df; font-size: 24px; }}
    h3 {{ margin-top: 24px; font-size: 19px; }}
    p, li {{ font-size: 16px; line-height: 1.82; }}
    ul {{ padding-left: 22px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 14px 0 22px; background: #fff; border: 1px solid #d7e2dd; }}
    th, td {{ border: 1px solid #d7e2dd; padding: 10px 11px; vertical-align: top; font-size: 14px; line-height: 1.65; }}
    th {{ background: #eaf4ef; text-align: left; }}
    @media (max-width: 760px) {{ .topbar-inner {{ align-items: flex-start; flex-direction: column; padding: 12px 14px; }} main {{ padding: 22px 14px 56px; }} table {{ display: block; overflow-x: auto; }} h1 {{ font-size: 27px; }} }}
  </style>
</head>
<body>
  <div class="topbar">
    <div class="topbar-inner">
      <div class="topbar-title">奥森预测与调整方案</div>
      <div class="actions">
        <a class="primary" href="/static/osen_prediction_adjustment_report_20260607.docx" download>下载 DOCX 报告</a>
        <a href="/static/osen_prediction_adjustment_report_basis_20260607.json" target="_blank" rel="noreferrer">查看依据链</a>
      </div>
    </div>
  </div>
  <main>{body}</main>
</body>
</html>
"""
    path.write_text(page, encoding="utf-8")


def main() -> int:
    basis = build_basis()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QUALITY_DIR.mkdir(parents=True, exist_ok=True)
    STATIC_DIR.mkdir(parents=True, exist_ok=True)
    basis_path = QUALITY_DIR / "osen_prediction_adjustment_report_basis_20260607.json"
    md_path = OUT_DIR / "osen_prediction_adjustment_report_20260607.md"
    docx_path = OUT_DIR / "osen_prediction_adjustment_report_20260607.docx"
    html_path = OUT_DIR / "osen_prediction_adjustment_report_20260607.html"
    static_html_path = STATIC_DIR / "osen_prediction_adjustment_report_20260607.html"

    md = report_markdown(basis)
    basis_path.write_text(json.dumps(basis, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    md_path.write_text(md, encoding="utf-8")
    write_docx(basis, docx_path)
    write_html(md, html_path)
    write_html(md, static_html_path)
    result = {
        "status": "pass",
        "basis": str(basis_path),
        "markdown": str(md_path),
        "docx": str(docx_path),
        "html": str(html_path),
        "static_html": str(static_html_path),
        "docx_bytes": docx_path.stat().st_size,
        "html_bytes": html_path.stat().st_size,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
